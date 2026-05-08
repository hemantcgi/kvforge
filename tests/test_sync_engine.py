import sys, sqlite3
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))


def test_sync_db_created_on_init(tmp_path):
    from core.sync_engine import SyncStateDB
    db = SyncStateDB(str(tmp_path / "sync.db"))
    assert (tmp_path / "sync.db").exists()


def test_sync_db_tables_exist(tmp_path):
    from core.sync_engine import SyncStateDB
    db = SyncStateDB(str(tmp_path / "sync.db"))
    conn = sqlite3.connect(str(tmp_path / "sync.db"))
    tables = {r[0] for r in conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()}
    conn.close()
    assert "document_hashes" in tables
    assert "section_hashes" in tables
    assert "deleted_docs" in tables
    assert "sync_runs" in tables


def test_upsert_and_get_doc_hash(tmp_path):
    from core.sync_engine import SyncStateDB
    db = SyncStateDB(str(tmp_path / "sync.db"))
    db.upsert_doc_hash("uc1", "file_001", "abc123", "2026-01-01T00:00:00+00:00")
    row = db.get_doc_hash("uc1", "file_001")
    assert row["doc_hash"] == "abc123"


def test_upsert_and_get_section_hash(tmp_path):
    from core.sync_engine import SyncStateDB
    db = SyncStateDB(str(tmp_path / "sync.db"))
    db.upsert_section_hash("uc1", "file_001", "slide:1", "deadbeef", '["chunk_a"]', "2026-01-01T00:00:00+00:00")
    rows = db.get_section_hashes("uc1", "file_001")
    assert rows[0]["section_id"] == "slide:1"
    assert rows[0]["content_hash"] == "deadbeef"


def test_list_source_ids(tmp_path):
    from core.sync_engine import SyncStateDB
    db = SyncStateDB(str(tmp_path / "sync.db"))
    db.upsert_doc_hash("uc1", "file_a", "h1", "2026-01-01T00:00:00+00:00")
    db.upsert_doc_hash("uc1", "file_b", "h2", "2026-01-01T00:00:00+00:00")
    ids = db.list_source_ids("uc1")
    assert "file_a" in ids
    assert "file_b" in ids


def test_record_and_get_sync_runs(tmp_path):
    from core.sync_engine import SyncStateDB
    db = SyncStateDB(str(tmp_path / "sync.db"))
    db.record_sync_run(
        "uc1",
        started_at="2026-05-01T10:00:00+00:00",
        finished_at="2026-05-01T10:01:00+00:00",
        files_checked=100,
        files_changed=5,
        chunks_added=12,
        chunks_superseded=3,
        pii_detections=0,
        errors="",
    )
    runs = db.get_sync_runs("uc1")
    assert len(runs) == 1
    assert runs[0]["files_checked"] == 100
    assert runs[0]["finished_at"] == "2026-05-01T10:01:00+00:00"


def test_sync_runs_newest_first(tmp_path):
    from core.sync_engine import SyncStateDB
    db = SyncStateDB(str(tmp_path / "sync.db"))
    db.record_sync_run("uc1", started_at="2026-01-01T00:00:00+00:00")
    db.record_sync_run("uc1", started_at="2026-01-02T00:00:00+00:00")
    runs = db.get_sync_runs("uc1")
    assert runs[0]["started_at"] == "2026-01-02T00:00:00+00:00"


def test_sync_runs_uc_isolation(tmp_path):
    from core.sync_engine import SyncStateDB
    db = SyncStateDB(str(tmp_path / "sync.db"))
    db.record_sync_run("uc1", started_at="2026-01-01T00:00:00+00:00")
    db.record_sync_run("uc2", started_at="2026-01-02T00:00:00+00:00")
    assert len(db.get_sync_runs("uc1")) == 1
    assert len(db.get_sync_runs("uc2")) == 1


import hashlib, json, tempfile
from datetime import datetime, timezone
from unittest.mock import MagicMock, patch
from docx import Document


def _make_docx_bytes(paragraphs: list[tuple]) -> bytes:
    import io
    doc = Document()
    for style, text in paragraphs:
        if style.startswith("Heading"):
            doc.add_heading(text, level=int(style.split()[-1]))
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _mock_connector(files: list[dict]) -> MagicMock:
    from connectors.base import SourceFile
    conn = MagicMock()
    conn.supports_delta.return_value = False
    source_files = [
        SourceFile(
            id=f["id"], name=f["name"], path=f["name"],
            size=len(f["content"]),
            modified_at=datetime(2026, 1, 1, tzinfo=timezone.utc),
        )
        for f in files
    ]
    conn.list_files.return_value = source_files
    content_map = {f["id"]: f["content"] for f in files}
    conn.download.side_effect = lambda sf: content_map[sf.id]
    return conn


def test_sync_engine_indexes_new_file(tmp_path):
    from core.sync_engine import SyncEngine, SyncStateDB
    db = SyncStateDB(str(tmp_path / "sync.db"))
    content = _make_docx_bytes([("Heading 1", "Intro"), ("Normal", "Body text.")])
    connector = _mock_connector([{"id": "doc1", "name": "doc1.docx", "content": content}])

    indexed_chunks = []
    mock_indexer = MagicMock(side_effect=lambda cfg, chunks: indexed_chunks.extend(chunks))

    engine = SyncEngine(uc_name="uc1", db=db, connector=connector, cfg=MagicMock(), indexer=mock_indexer)
    stats = engine.run()

    assert stats["files_changed"] == 1
    assert len(indexed_chunks) >= 1


def test_sync_engine_skips_unchanged_file(tmp_path):
    from core.sync_engine import SyncEngine, SyncStateDB
    db = SyncStateDB(str(tmp_path / "sync.db"))
    content = _make_docx_bytes([("Normal", "Static content.")])
    connector = _mock_connector([{"id": "doc1", "name": "doc1.docx", "content": content}])

    indexed_chunks = []
    mock_indexer = MagicMock(side_effect=lambda cfg, chunks: indexed_chunks.extend(chunks))

    engine = SyncEngine(uc_name="uc1", db=db, connector=connector, cfg=MagicMock(), indexer=mock_indexer)
    engine.run()  # first run — indexes everything
    indexed_chunks.clear()
    engine.run()  # second run — same content, nothing should be re-indexed

    assert len(indexed_chunks) == 0


def test_sync_engine_detects_deletion(tmp_path):
    from core.sync_engine import SyncEngine, SyncStateDB
    db = SyncStateDB(str(tmp_path / "sync.db"))
    content = _make_docx_bytes([("Normal", "Will be deleted.")])

    # First run with the file present
    connector = _mock_connector([{"id": "doc1", "name": "doc1.docx", "content": content}])
    engine = SyncEngine(uc_name="uc1", db=db, connector=connector, cfg=MagicMock(), indexer=MagicMock())
    engine.run()

    # Second run with the file gone
    connector2 = _mock_connector([])
    engine2 = SyncEngine(uc_name="uc1", db=db, connector=connector2, cfg=MagicMock(), indexer=MagicMock())
    stats2 = engine2.run()

    assert stats2["files_deleted"] >= 1
    remaining = db.list_source_ids("uc1")
    assert "doc1" not in remaining
