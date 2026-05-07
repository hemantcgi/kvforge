import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest
from docx import Document
from docx.shared import Pt


def _make_docx(tmp_path, paragraphs: list[tuple]) -> str:
    """paragraphs: list of (style, text). style is 'Normal', 'Heading 1', etc."""
    doc = Document()
    for style, text in paragraphs:
        if style.startswith("Heading"):
            level = int(style.split()[-1])
            doc.add_heading(text, level=level)
        else:
            doc.add_paragraph(text, style=style)
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return str(path)


def test_docx_loads_plain_text(tmp_path):
    path = _make_docx(tmp_path, [("Normal", "Hello world paragraph.")])
    from ingestion.docx_loader import DocxLoader
    chunks = DocxLoader().load(path)
    assert len(chunks) >= 1
    assert "Hello world paragraph." in chunks[0]["text"]


def test_docx_heading_metadata(tmp_path):
    path = _make_docx(tmp_path, [
        ("Heading 1", "Introduction"),
        ("Normal", "This is the intro body text."),
    ])
    from ingestion.docx_loader import DocxLoader
    chunks = DocxLoader().load(path)
    assert chunks[0]["metadata"]["heading_text"] == "Introduction"
    assert chunks[0]["metadata"]["heading_level"] == 1


def test_docx_nested_headings(tmp_path):
    path = _make_docx(tmp_path, [
        ("Heading 1", "Chapter One"),
        ("Heading 2", "Section 1.1"),
        ("Normal", "Content under 1.1."),
    ])
    from ingestion.docx_loader import DocxLoader
    chunks = DocxLoader().load(path)
    meta = {c["metadata"]["heading_text"]: c["metadata"]["heading_level"] for c in chunks}
    assert meta.get("Section 1.1") == 2


def test_docx_table_chunk(tmp_path):
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Score"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "90"
    table.cell(2, 0).text = "Bob"
    table.cell(2, 1).text = "85"
    path = tmp_path / "table.docx"
    doc.save(str(path))
    from ingestion.docx_loader import DocxLoader
    chunks = DocxLoader().load(str(path))
    texts = " ".join(c["text"] for c in chunks)
    assert "Alice" in texts
    assert "Name" in texts  # headers prepended


def test_docx_source_file_metadata(tmp_path):
    path = _make_docx(tmp_path, [("Normal", "Some text.")])
    from ingestion.docx_loader import DocxLoader
    chunks = DocxLoader().load(path)
    assert chunks[0]["metadata"]["source"] == Path(path).name


def test_docx_chunk_id_present(tmp_path):
    path = _make_docx(tmp_path, [
        ("Heading 1", "Chapter"),
        ("Normal", "Text one. Text two."),
    ])
    from ingestion.docx_loader import DocxLoader
    chunks = DocxLoader().load(path)
    # Every chunk should have a chunk_id
    assert all("chunk_id" in c["metadata"] for c in chunks)
    # chunk_id should be 0-indexed and increment
    for i, chunk in enumerate(chunks):
        assert chunk["metadata"]["chunk_id"] == i


def test_docx_section_hash_present(tmp_path):
    path = _make_docx(tmp_path, [
        ("Heading 1", "Intro"),
        ("Normal", "Body text."),
    ])
    from ingestion.docx_loader import DocxLoader
    chunks = DocxLoader().load(path)
    assert "section_hash" in chunks[0]["metadata"]
    assert len(chunks[0]["metadata"]["section_hash"]) == 64  # SHA-256 hex


def test_docx_empty_file(tmp_path):
    doc = Document()
    path = tmp_path / "empty.docx"
    doc.save(str(path))
    from ingestion.docx_loader import DocxLoader
    chunks = DocxLoader().load(str(path))
    assert chunks == []


def test_docx_table_chunk_metadata(tmp_path):
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Score"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "90"
    table.cell(2, 0).text = "Bob"
    table.cell(2, 1).text = "85"
    path = tmp_path / "table_meta.docx"
    doc.save(str(path))
    from ingestion.docx_loader import DocxLoader
    chunks = DocxLoader().load(str(path))
    assert all(c["metadata"]["is_table"] for c in chunks)
    assert all("table_position" in c["metadata"] for c in chunks)
    assert chunks[0]["metadata"]["table_position"]["row"] == 1
    assert chunks[1]["metadata"]["table_position"]["row"] == 2


def test_docx_table_row_format(tmp_path):
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Score"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "90"
    path = tmp_path / "table_fmt.docx"
    doc.save(str(path))
    from ingestion.docx_loader import DocxLoader
    chunks = DocxLoader().load(str(path))
    assert len(chunks) == 1
    assert "Name: Alice" in chunks[0]["text"]
    assert "Score: 90" in chunks[0]["text"]


def test_docx_chunk_size_and_overlap(tmp_path):
    # 120 words so chunk_size=50 should produce at least 2 chunks
    words = " ".join(f"word{i}" for i in range(120))
    path = _make_docx(tmp_path, [("Heading 1", "Big Section"), ("Normal", words)])
    from ingestion.docx_loader import DocxLoader
    chunks = DocxLoader(chunk_size=50, chunk_overlap=10).load(path)
    # Should produce multiple chunks from the one long section
    assert len(chunks) >= 2
    # Second chunk should start with words from the overlap region
    first_last_words = chunks[0]["text"].split()[-10:]
    second_first_words = chunks[1]["text"].split()[:10]
    # There should be overlap between end of chunk 0 and start of chunk 1
    assert any(w in second_first_words for w in first_last_words)


def test_docx_same_section_hash_across_windows(tmp_path):
    words = " ".join(f"word{i}" for i in range(120))
    path = _make_docx(tmp_path, [("Heading 1", "Big Section"), ("Normal", words)])
    from ingestion.docx_loader import DocxLoader
    chunks = DocxLoader(chunk_size=50, chunk_overlap=10).load(path)
    assert len(chunks) >= 2
    # All chunks from the same section must share the same section_hash
    hashes = {c["metadata"]["section_hash"] for c in chunks}
    assert len(hashes) == 1
